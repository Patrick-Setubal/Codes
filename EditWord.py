
# --------------------------Importações de Bibliotecas-------------------------
from docx import Document # Trabalhar com Word
import python_docx_replace as py_replace #(Alterar para ganho de desempenho)


# import re
# from python_docx_replace.paragraph import Paragraph

# __all__ = ["docx_replace", "docx_blocks", "docx_remove_table"]

# def docx_replace(doc, **kwargs: str) -> None:
#     '''
#     Replace all the keys in the word document with the values in the kwargs

#     ATTENTION: The required format for the keys inside the Word document is: ${key}

#     Example usage:
#         Word content = "Hello ${name}, your phone is ${phone}, is that okay?"

#         doc = Document("document.docx")  # python-docx dependency

#         docx_replace(doc, name="Ivan", phone="+55123456789")

#     More information: https://github.com/ivanbicalho/python-docx-replace
#     '''
    
#     for p in Paragraph.get_all(doc):
#         paragraph = Paragraph(p)
#         if "$" in paragraph.p.text:
#             find_Keys = re.findall(r'\$\{[\w*\-*\.*]*\}', paragraph.p.text)
#             for find_key in find_Keys:
#                 kwargs[find_key[2:-1]]
#                 paragraph.replace_key(find_key, str(kwargs[find_key[2:-1]]))






from docx.shared import Inches  # Trabalhar com parametros de imagem do docx
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Formatar Paragrafos
import os

import re #Função para localizar strings 

# --------------------------Configurações de Variaveis-------------------------
# Local e texto padrão de onde estão os modelos prontos
Modelos = r'C:/Users/Colaborador/Desktop/name_Coleta-3/Coleta App/'

# --------------------------Automações-----------------------------------------
# Função Para Utilizar


def GerarRelatorio(CtrAut, coleta) -> str:
    '''
    Objetivo:
        Gerar Relatorio

    Parametros:
        CtrAut: Controle da automação
        coleta: Coleta pendente

    Retorno:[
        Nome do Processo
        Status Do Processo
        ErroRelatorio
        Outros ...
    ]
    '''

    # ---------------------- Variaveis ----------------------
    Adress_folder = CtrAut['adress']['adress_folder']
    name_coleta = CtrAut['name_coleta']
    try:
        try:
            
            # Pegar Documento Modelo
            doc = Document(Adress_folder+"/"+name_coleta+"/Relatorios/Relatorio_Modelo_"+name_coleta+".docx")
            
            # Definir destino do relatorio
            Adress_Doc_Result = Adress_folder+"/"+name_coleta+"/Relatorios/"+coleta['TAG']+" - "+str(coleta['ID'])+".docx"
            
            # Pegar endereço das Imagens
            Adress_Img = Adress_folder+"/"+name_coleta+"/Imagens/"
            
            # Limpando valores na
            coleta = coleta.fillna('-')
            
            # Trasnformar DF da coleta em dict
            Inputs = dict(coleta)
            
        except Exception as e:
            CtrAut['Aut']['GerarRelatorio']['Lista_Erros'].append(
                f"Erro Geral das variaveis: {str(e)}")
            

        # ---------------------- Codigo ----------------------

        #Preencher Variaveis do relatorio 
        try:
            # Editar todas as variaveis do Documento
            # Utilizar Bibliooteca "python_docx_replace"
            # Com algumas modificações para ganho de performance
            py_replace.docx_replace(doc, **Inputs)
        except Exception as e:
            CtrAut['Aut']['GerarRelatorio']['Lista_Erros'].append(
                f"Erro Geral do preenchimento do Relatorio: {str(e)}")
        
        #Inserir imagens no Relatorio 
        try:
            # Fluxo para Adicionar Imagem
            remove_rows = []
            for table in doc.tables:
                for row in table.columns:
                    for i, cell in enumerate(row.cells):
                        if "$Imagem" in cell.text:
                            ModelInputs = re.findall(r'\$[\w*\-*\.*]*\$', cell.text)
                            # Percorrer todos os $Imagem..$
                            for ModelInput in ModelInputs:
                                # Se for Para Adicionar Imagem
                                if CtrAut['Aut']['GerarRelatorio']['Exec_img']:
                                    # Procurar nome da imagem
                                    Nome_Imagem = str(
                                        coleta[(ModelInput.replace("$", ""))])

                                    # Se So tiver o codigo padrão ou vazio, deletar linha
                                    if len(Nome_Imagem) == 16 or Nome_Imagem == "":
                                        remove_rows.insert(0,[table,i]) 
                                    else:
                                        # Separo endereço da Imagem
                                        Img = Adress_Img+Nome_Imagem+".png"
                                        # Limpar conteudo Escrito
                                        cell.text = cell.text.replace(ModelInput, "")
                                        # Adicionar imagem e Centralizar
                                        try:
                                            # Adicionar Imagem
                                            cell.add_paragraph().add_run().add_picture(Img, width=Inches(4.0))
                                            # Centralizar Imagem
                                            cell.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            # Adicionar Nome da Imagem
                                            cell.add_paragraph().add_run(coleta[(ModelInput.replace("$", ""))][16:])
                                            # Centralizar nome da imagem 
                                            cell.paragraphs[2].alignment = WD_ALIGN_PARAGRAPH.CENTER

                                        except Exception as e:
                                            CtrAut['Aut']['GerarRelatorio']['Lista_Erros'].append(
                                                f"Erro na Imagem: {str(e)}")
                                            cell.text = "Erro"
                                # Se nao for para adicionar Imagem
                                else:
                                    # Limpar conteudo Escrito
                                    remove_rows.insert(0,[table,i]) 

            # Deletar slite de imagens vazios 
            for remove_row in remove_rows:
                remove_row[0]._tbl.remove(remove_row[0].rows[remove_row[1] ]._tr)
        except Exception as e:
            CtrAut['Aut']['GerarRelatorio']['Lista_Erros'].append(
                f"Erro Geral da Imagem: {str(e)}")
        
        # Salvar Documento
        try:
            doc.save(Adress_Doc_Result)
        except Exception as e:
            CtrAut['Aut']['GerarRelatorio']['Lista_Erros'].append(
                f"Erro Geral salvamento do relatorio: {str(e)}")
        
        # # Criar lista de endereço das imagem (Usado quando quiser deletar as imagens)
        # Adress_img = []
        # for i in range(5):
        #     coluna_imagem = 'Imagem'+str(int((i+1)))
        #     Nome_Imagem = coleta[coluna_imagem]
        #     if len(Nome_Imagem) != 16:
        #         if Nome_Imagem != "":
        #             Adress_img.append(os.path.abspath(Adress_Img+Nome_Imagem+".png"))
                
        
        # Verificar se houve erro e atualizar status
        if len(CtrAut['Aut']['GerarRelatorio']['Lista_Erros']) == 0:
            CtrAut['Aut']['GerarRelatorio']['status'] = 'OK'
        else:
            CtrAut['Aut']['GerarRelatorio']['status'] = 'Erro'
                        
        # Salvar endereço do Relatorio
        CtrAut['Aut']['GerarRelatorio']['adress_relatorio'] = os.path.abspath(Adress_Doc_Result)
        
        return CtrAut
    
    
    except Exception as e:
        CtrAut['Aut']['GerarRelatorio']['Lista_Erros'].append(
            f"Erro Geral na Automação Relatorio: str{e}")
        
        CtrAut['Aut']['GerarRelatorio']['status'] = 'Erro'
        return CtrAut
    